# password.txt 第一行 id 第二行秘钥 第三行fromlang 第四行tolang

import http.client as httplib
import hashlib
import urllib.parse
import random
import json
class trans():
    def __init__(self,id,sk,fromLang,toLang):
        self.appKey = id
        self.secretKey = sk
        self.fromLang = fromLang
        self.toLang = toLang
        self.md5 = hashlib.md5()
    def trans(self,q):
        httpClient = None
        myurl = '/api'
        salt = random.randint(1, 65536)
        sign = self.appKey + q + str(salt) + self.secretKey
        self.md5.update(sign.encode("utf8"))
        sign = self.md5.hexdigest()
        myurl = myurl + '?appKey=' + self.appKey + '&q=' + urllib.parse.quote(q) + '&from=' + self.fromLang + '&to=' + self.toLang + '&salt=' + str(
            salt) + '&sign=' + sign
        httpClient = httplib.HTTPConnection('openapi.youdao.com')
        httpClient.request('GET', myurl)
        response = httpClient.getresponse()
        a= response.read()
        return json.loads(a,encoding='utf8')


password = open('./password.txt')
lines = password.readlines()
id = lines[0].strip()
secretKey = lines[1].strip()
fromLang = lines[2].strip()
toLang = lines[3].strip()

import os
docs = [word for word in os.listdir("./") if ".docx" in word]


import docx
for doc in docs:
    file=docx.Document("./%s"%doc)
    file2 = docx.Document()
    file2_list=[]
    for para in file.paragraphs:
        print(para.text)
        file2.add_paragraph(para.text)
        tran = trans(id=id, sk=secretKey, fromLang=fromLang, toLang=toLang)
        test= tran.trans(para.text)
        try:
            print(test["translation"][0])
            file2_list.append(test["translation"][0])
        except:
            pass
    file2.add_paragraph("\n")
    for para in file2_list:
        file2.add_paragraph(para)
    file2.save('trans-%s'%doc)